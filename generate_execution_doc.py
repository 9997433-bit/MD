#!/usr/bin/env python3
"""Generate narrative-style thesis execution guide (minimal tables)."""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATHS = [
    Path("/workspace/直驱三轴龙门动态模型确认_逐步执行总表.docx"),
    Path("/workspace/thesis_execution_table.docx"),
    Path("/workspace/浙大/毕业设计/仿真-实验大纲/直驱三轴龙门动态模型确认_逐步执行总表.docx"),
    Path("/workspace/download/直驱三轴龙门动态模型确认_逐步执行总表.docx"),
]


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.space_after = Pt(6)


def add_para(doc: Document, text: str, style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def clean_body(body: str) -> str:
    """Remove inline 必存 lines; deliverables are rendered in dedicated sections."""
    paras = []
    for p in body.strip().split("\n\n"):
        p = p.strip()
        if not p:
            continue
        if re.match(r"^必存[：:]", p):
            continue
        # drop trailing sentence chunks that are only deliverable reminders
        p = re.sub(r"\n?必存[文件数据图片][：:][^\n]+$", "", p).strip()
        if p:
            paras.append(p)
    return "\n\n".join(paras)


def add_step(
    doc: Document,
    code: str,
    title: str,
    body: str,
    checkpoints: list[str] | None = None,
    deliverables: dict | None = None,
) -> None:
    doc.add_paragraph(f"{code}　{title}", style="Heading 3")
    for para in clean_body(body).split("\n\n"):
        add_para(doc, para.strip())
    if deliverables:
        data_items = deliverables.get("data") or []
        image_items = deliverables.get("images") or []
        add_para(doc, "【必存数据】（本步完成后必须存档的文件/文件夹，命名保持一致）")
        if data_items:
            for item in data_items:
                doc.add_paragraph(item, style="List Number")
        else:
            add_para(doc, "本步无单独数据文件；操作记录写入 Experiment_Matrix 或对应阶段日志即可。")
        add_para(doc, "【必存图片】（本步必须拍摄或导出的图，用于论文/答辩/追溯）")
        if image_items:
            for item in image_items:
                doc.add_paragraph(item, style="List Number")
        else:
            add_para(doc, "本步无专用图片要求。")
    if checkpoints:
        add_para(doc, "现场操作要点（逐项打勾）：")
        for item in checkpoints:
            doc.add_paragraph(item, style="List Bullet")


STEPS = {
    "阶段0：实验前准备（必须先于一切实验）": [
        (
            "0-1",
            "建立文件夹体系与 Experiment_Matrix（实验矩阵总表）",
            """在本步骤中，你需要为后续全部实验（B1 区 A 扫描、B2 Y 极限扫描、B3 盲验、谐响应等）预先建立层级清晰的文件夹，并编制 Experiment_Matrix 总表。推荐标准目录树如下（可微调命名，但层级逻辑不可乱）：

01_Admin/ — 矩阵、预约、安全告知、Pre-registration、门禁签字 PDF  
02_Test_Raw/ — 全部实验原始数据（B1/B2/B3/谐响应，只读封存）  
03_Simulation/ — Baseline_FE_v0、Updated_FE_ThetaStar、FE_Y_Extreme、谐响应仿真结果  
04_SCI_Reserve/ — 10-start、EMA 汇总、5 trial 完整备份（与 θ* 版本隔离）  
05_Thesis/ — 论文章节、Figure_Index、附录、Public 脱敏版  

Experiment_Matrix 总表至少应包含：实验编号、执行日期、机床轴位、SSS 版本（SSS-M 或 SSS-H）、激励点、操作者、原始数据路径、备份路径、是否只读封存。每个实验编号对应唯一路径，命名建议采用「Exp编号_构型_方法_Raw」形式（例如 ExpB1_C0_SLDV_Raw），避免仅用日期或「试验1」这类无法追溯的名称。

这一步骤的核心价值在于：从数据产生的第一分钟起就保证可追溯性。若文件夹命名随意、或实验结束后才补建目录，答辩时很难证明「这些数据确实来自同一次、同一状态下的测量」。常见遗漏是只建了物理文件夹却没有 Experiment_Matrix，或矩阵里缺少备份路径与操作者字段。Fatal 级风险是数据散乱、无法与 SSS（标准状态）记录、轴位截图一一对应。

完成后必须保存 Experiment_Matrix.xlsx；本步骤无需额外图片，但建议在总表中增加一列「关联图片文件名」以便后续索引。""",
        ),
        (
            "0-2",
            "P10–P12 事前登记（Pre-registration）并获导师签字",
            """盲验（Blind Validation）能否成立，取决于 P10、P11、P12 三个点的坐标是否在正式实验与模型修正之前就已被固定，且与修正过程完全隔离。本步骤要求：在一切 SLDV/EMA 正式测量之前，书面登记三点的空间坐标、测点类型（结合部附近 / 刀位旁 / 空间远端）、对应实物标记方式，并由导师签字确认。三点应在结构上空间分散，分别代表不同物理区域的响应特征，不可挤在同一条梁面上。

若事后根据修正效果「挑选」盲验点，则盲验在学术逻辑上不成立，属于 Fatal 级问题。常见遗漏是只登记 P10 而忽略 P11、P12，或登记坐标与实物标记不一致（例如图纸坐标与现场贴点位置偏差未记录）。注意：登记时间必须早于 B1 正式扫频，且登记文件应单独存放，修正阶段（D 阶段）严禁打开盲验原始数据文件夹。

必存文件：P10-P12_PreRegistration.pdf（含坐标、示意图、签字页）。必存图片：测点标记实物照片（需能看清点位编号与贴点位置）。""",
        ),
        (
            "0-3",
            "定稿 SSS（Standard State Specification，标准状态规范）及 SSS-M / SSS-H 双版本",
            """SSS 描述的是「仿真与实验必须共同遵守的机床状态」，包括但不限于：各直线轴位置、伺服上电/下电、主轴启停、冷却/风机、拖链与电缆姿态、工装夹具、地脚垫铁与拧紧状态等。对于直驱三轴龙门，务必明确直驱电机是否通电、锁轴状态如何——静止不等于状态一致。

本课题涉及两种构型，必须分别定稿两份状态规范：SSS-M（Y 中位构型，对应 B1 标定实验与 Baseline FE）和 SSS-H（Y 极限构型，对应 B2 实验与 FE_Y_Extreme）。两份文件除 Y 轴坐标、拖链/电缆姿态外，其余电气与辅机要求应保持一致并交叉引用。每次实验前用对应版本的 SSS_Checklist 逐项核对并签字；B2、B3 若与 B1 不同日进行，实验当日仍须重走 B0 环境/噪底/SSS 核对，不可引用 B1 当天的记录代替。

若仿真按某一轴位与通电状态建模，而实验时拖链下垂、电机下电或相邻设备未关，会导致频率系统性偏移，后续 MAC（Modal Assurance Criterion，模态置信准则）对比失去意义。常见遗漏是只写了几何轴位、未区分 SSS-M/SSS-H，或未写电气与辅机状态。易错点是 B2 实验仍用 SSS-M 检查表。""",
        ),
        (
            "0-4",
            "定稿 MAC 法向投影协议",
            """SLDV（Scanning Laser Doppler Vibrometry，扫描激光测振）测量的是沿激光束方向的振动分量，而有限元振型通常是三维向量。因此，在计算 MAC 之前，必须将 FE 振型投影到各测点的激光法向方向，再与实验振型做 MAC 对比。本步骤要求书面定稿：投影公式、归一化方式、MAC 阈值定义、多测点聚合规则（若适用），并与后续 C2 配对、D 阶段修正、E 阶段盲验全部保持一致。

若直接用三维振型向量算 MAC，或不同章节采用不同归一化方式，会出现「MAC 假低」或「修正方向相反」的 Fatal 风险。常见遗漏是没有单独成文的 MAC_Calculation_Protocol，导致论文第 3、4 章公式与实际操作不一致。本协议一旦定稿，全文不得随意更改；若必须更改，需记录版本号并说明对已有结果的影响。

必存文件：MAC_Calculation_Protocol.pdf。本步骤无需额外实验图片，但建议在协议中附一张法向投影示意图。""",
        ),
        (
            "0-5",
            "定稿 Trial 1–3 选用规则",
            """实验设计中每次测量通常做 5 次重复（Trial 1–5），但论文主结果与修正过程一般报告 Trial 1–3 的均值。本步骤要求事先书面规定：如何从 5 次重复中选取 Trial 1–3、何种情况下剔除某次 trial（例如力幅异常、相干系数 γ² 过低、明显操作失误）、剔除是否需导师确认。规则一旦定稿，全文统一引用，不得事后再挑「看起来最好」的三次。

若选用规则不明确，审稿人或答辩委员会会质疑结果是否被「挑选」。常见遗漏是只做了 5 次重复却没有成文的选用与剔除规则。易错点是实验时未逐 trial 填写 Trial_Log，导致事后无法解释为何剔除某次数据。

必存文件：Trial_Selection_Rule.pdf。建议在规则中附重复性判据（与阶段 0 末尾的合格指标一致）。""",
        ),
        (
            "0-6",
            "坐标变换定义（机床 / CAD / FE / SLDV）",
            """P1–P12、激励点 E1–E3、刀位点 TP 的坐标必须在机床坐标系、CAD 坐标系、有限元坐标系、SLDV 测量坐标系之间建立明确变换关系，单位统一为 mm，正方向与右手系约定写清。任何一步坐标系不一致，都会导致测点映射偏差，进而使 MAC 系统性偏低，修正算法向错误方向收敛。

本步骤应产出 Coordinate_Transform.pdf，包含：各坐标系原点定义、变换矩阵或对齐步骤、验证方法（例如用 P1 试测验证映射距离 d）。常见遗漏是 CAD 与 FE 原点不同但未文档化，或 SLDV 软件内坐标与 FE 导出坐标未做一次性核对。Fatal 风险是映射距离 d 大于 10 mm 却未复核仍用于修正。

必存文件：Coordinate_Transform.pdf。必存图片：坐标系关系示意图（建议在 CAD 或 FE 中标注原点与轴向）。""",
        ),
        (
            "0-7",
            "模态筛选规则（剔除刚体 / 局部，保留 6 阶弹性模态）",
            """仿真与实验均可能提取较多模态阶次（建议提取 12 阶），但论文与修正过程以 6 阶弹性模态为主。本步骤要求书面定义：如何识别并剔除刚体模态、局部模态或数值假模态，保留哪 6 阶作为「弹性模态」参与配对与修正。判据可包括：频率范围、振型能量分布、与预期物理形态的一致性、稳态图（Stabilization Diagram）上的稳定极点等。

若筛选规则不清晰，会出现「按阶号硬对」——实验第 3 阶其实对应 FE 第 5 阶——导致越修越偏。常见遗漏是剔除过程没有单独记录，论文中只给出 6 阶结果却无法解释另外 6 阶为何不用。易错点是把明显局部振动当作整体模态参与 MAC 优化。

必存文件：Mode_Screening_Rules.docx。必存图片：被剔除模态与保留模态的振型对比示例（至少 1 组）。""",
        ),
        (
            "0-8",
            "预约机床、SLDV 与备份实验时段",
            """B1（Y 中位 + 区 A）、B2（Y 极限 raw）、B3（P10–P12 盲验）应在时间上连续或紧凑完成，避免间隔过长导致机床状态、环境温度、操作人员发生变化。本步骤完成设备预约、人员协调，并预留 B4 备用重测时段——若 B1–B3 任一步失败，仅重测失败步骤，而非全盘重来。

常见遗漏是没有预留备用实验窗口，导致正式数据不合格时只能挤占其他课题时间或接受不合格数据。易错点是把「摸底试验」与「正式试验」混为一次，事后无法证明正式数据是在 G1/G2 门禁通过之后采集的。

必存文件：预约记录（邮件或实验室系统截图均可）。本步骤无需额外图片。""",
        ),
        (
            "0-9",
            "采购耗材、安全告知与激光防护",
            """确认反光膜、加速计、力锤、力传感器、连接线缆等耗材到位；对参与实验人员完成激光安全告知（SLDV  Class 等级、佩戴防护、禁止直视等），安全签字存档。缺反光膜或测点准备不足会导致 SLDV 信噪比（SNR）不够，数据不可用。

常见遗漏是安全告知未签字，或正式实验时才发现反光膜不足临时补贴导致测点位置偏移。本步骤完成后，检查清单上所有项应可在 B0 实验前再次快速核对。

必存文件：采购清单、安全告知签字页。本步骤无需额外图片。""",
        ),
    ],
    "阶段A：Baseline 仿真（基准有限元）": {
        "A1 几何与质量": [
            (
                "A1-1",
                "导入 CAD 并统一坐标系",
                """将机床 CAD 装配体导入有限元前处理环境，完成几何清理与简化，并确保 FE 模型坐标系与 Coordinate_Transform.pdf 完全一致。所有简化（删小孔、合并面、忽略非结构件等）必须写入简化清单，说明对质量和刚度可能的影响方向，不可静默删除筋板、电机壳体或拖链等可能影响低阶模态的部件。

若坐标系不一致，后续 P1–P12 全部映射错误，属于 Fatal 级问题。常见遗漏是未记录简化清单，答辩时无法回答「为什么比 CAD 轻/软」。易错点是漏装直线电机、拖链或主轴组件质量，导致频率整体偏高。

必存数据：A1_Model.step（或同等 CAD/几何源文件）。必存图片：CAD 装配图（标注主要部件与坐标原点）。""",
            ),
            (
                "A1-2",
                "质量与质心核对（FE vs BOM / 称重）",
                """将 FE 模型总质量、各部件质量、质心位置与 BOM（物料清单）或实际称重结果对比。合格目标：总质量误差小于 5%（优秀小于 3%）。若误差超标，应优先检查简化过度、材料密度错误或漏装部件，而不是在后续修正阶段用「参数微调」掩盖质量错误。

质量误差会直接造成频率系统性偏差，使结合部参数失去物理意义。常见遗漏是没有 Mass_Check 记录表。易错点是为了网格好划分而删除筋板或立柱局部加强结构，导致刚度与质量同时偏离。

必存数据：Mass_Summary.csv。必存图片：质心位置示意图（FE 与参考值对比）。""",
            ),
        ],
        "A2 结合部与边界": [
            (
                "A2-1",
                "结合部参数化 θ₀（导轨 / 电机座 / 立柱等）",
                """对导轨连接、电机座、立柱结合面等难以精确建模的位置，引入 6–8 个有物理依据的参数 θ₀（例如接触刚度、阻尼比例、螺栓等效刚度），写入 Joint_Theta0.csv，每个参数给出初值、合理上下界（Bounds）及文献/经验来源。参数过少无法反映结合部柔度，过多则修正不可识别且易过拟合。

Fatal 风险是全刚性连接——低阶模态频率往往显著偏高，与实验完全无法配对。易错点是一次性引入超过 10 个参数，或参数无 Bounds 约束导致优化跑出负刚度。

必存数据：Joint_Theta0.csv。必存图片：结合部位置示意图（标注参数对应物理位置）。""",
            ),
            (
                "A2-2",
                "地脚边界条件（BC）与实验状态一致",
                """地脚垫铁、螺栓预紧、是否全约束等边界条件必须与 SSS 及实际安装一致。若实验机床并非「地脚全固定」，而 FE 采用全固定边界，低阶模态尤其是整体摆动模态会对不上。重力方向、约束方向务必核对，单位统一（N、mm、MPa）。

常见遗漏是 Boundary_Spec 未单独成文，或仿真与实验垫铁状态不一致。易错点是重力方向反了或约束施加在错误面上。

必存数据：Boundary_Spec.pdf。必存图片：边界条件标注图（FE 模型上标注约束位置与类型）。""",
            ),
        ],
        "A3 网格收敛": [
            (
                "A3-1",
                "粗 / 中 / 细三套网格收敛分析",
                """至少建立粗、中、细三套网格，对比前 3 阶弹性模态频率变化。当前 3 阶频率相对变化均小于 2% 时，可将中等或细网格定为 Baseline 网格，并记录于 Mesh_Convergence.csv。结合部、导轨接触区域应重点加密；这些区域网格过粗会产生虚假模态或错误振型。

跳过网格收敛是常见 Fatal 遗漏——审稿人几乎必问。易错点是只在整体结构上细化网格，而结合部仍然过粗，导致收敛曲线「看起来收敛了」但物理上不可信。

必存数据：Mesh_Convergence.csv。必存图片：网格划分图 + 前三阶频率收敛曲线。""",
            ),
        ],
        "A4 重力静力 + 预应力模态": [
            (
                "A4-1",
                "Static 重力静力分析",
                """在正确边界与材料参数下施加重力载荷，完成静力分析，检查变形趋势是否合理（例如龙门横梁下挠、立柱压缩等）。此步骤是 sanity check（合理性检查）：若静力变形方向或量级明显不合理，应先排查边界、单位与载荷，再进入模态分析。

常见遗漏是没有保存 Static 结果文件，后续无法证明预应力模态确实基于平衡态。易错点是单位混用（mm 与 m、MPa 与 Pa）导致变形量级错误却未察觉。

必存数据：Static_Results.rst（或同等结果文件）。必存图片：重力变形云图。""",
            ),
            (
                "A4-2",
                "Pre-stressed Modal（预应力模态）分析",
                """模态分析必须基于 Static 平衡态的预应力状态进行，不可对无预应力模型直接做模态并声称代表加工状态。记录分析设置：提取模态数、频率范围、是否考虑预应力效应等。与普通模态的结果差异应在论文中简要说明。

Fatal 风险是把普通模态与预应力模态混用，或只提取 6 阶而未保留足够阶次供后续筛选。易错点是静力未收敛就进入模态，或预应力步与模态步未正确链接。

必存数据：Modal_Results.rst。本步骤可复用后续振型图，不强制单独图片。""",
            ),
            (
                "A4-3",
                "提取 12 阶并筛选 6 阶弹性模态",
                """从预应力模态结果中提取至少 12 阶，再按 Mode_Screening_Rules 剔除刚体与局部模态，保留 6 阶弹性模态用于后续配对与展示。每一阶被剔除的模态都应记录剔除原因（频率过低、振型局部化等），写入 Elastic_Modes_6.csv。

常见遗漏是无剔除记录，答辩时无法解释为何选这 6 阶。易错点是把刚体模态或数值假模态当作弹性模态与实验配对。

必存数据：Elastic_Modes_6.csv。必存图片：Mode 1–6 振型图（FE）。""",
            ),
        ],
        "A5 映射与灵敏度": [
            (
                "A5-1",
                "P1–P12、E1–E3、TP 映射至 FE 节点",
                """将全部测点与刀位点 TP 映射到 FE 模型最近节点或插值位置，输出 P1-P12_Mapping.csv，记录映射距离 d。若 d 大于 10 mm，必须复核坐标变换或 CAD 对齐，不可强行用于 MAC 计算。P11、P12 等盲验点同样必须在 Baseline 阶段完成映射，但盲验数据在 D 阶段不得进入目标函数。

遗漏 P11/P12 映射是 Fatal 级错误——盲验将无法进行。易错点是 SLDV 坐标系与 FE 坐标系未统一，或 TP 刀位点未映射导致谐响应输出位置错误。

必存数据：P1-P12_Mapping.csv。必存图片：测点布置图（FE 与实验对照）。""",
            ),
            (
                "A5-2",
                "灵敏度分析与待识别参数确定",
                """对 θ₀ 各参数做灵敏度分析（频率或 MAC 对参数的偏导/有限差分），识别对目标模态影响最大的参数，优先选 4 个参数做试修正，再视 L-curve 等准则扩展至 6–8 个。灵敏度结果应支撑 Bounds 设置，避免修正阶段参数不可识别。

常见遗漏是没有 Sensitivity 分析就直接开始优化。易错点是一次性同时修正 8 个以上参数，导致多解性与过拟合。

必存数据：Sensitivity.csv。必存图片：灵敏度龙卷风图（Tornado Chart）。""",
            ),
        ],
        "A6 第二构型（若做 C1 构型验证）": [
            (
                "A6-1",
                "建立 FE_Y_Extreme（Y 轴极限构型模型）",
                """若论文包含 C1（Y 极限构型）验证，必须单独建立 Y 极限轴位下的 FE 模型 FE_Y_Extreme，拖链、电缆、平衡重等姿态与 B2 实验一致。θ* 冻结后对该构型做正推预测，禁止用 B2 的 raw 数据回头 refit θ*——否则构型验证失去独立意义。

Fatal 逻辑错误是用 Y 中位 Baseline FE 去对比 Y 极限实验。易错点是以为 θ* 会自动「适应」不同轴位——结合部参数可迁移，但几何构型必须更新。

必存数据：FE_Y_Extreme.wbpz（或项目文件）。本步骤不强制单独图片，建议在 SSS 中附 Y 极限轴位截图。""",
            ),
        ],
        "A7 门禁 G1": [
            (
                "A7-1",
                "冻结 Baseline_FE_v0 并通过 G1 门禁",
                """当 A1–A6 全部完成且自检合格后，将 Baseline 模型冻结为 Baseline_FE_v0，填写 G1 门禁记录并签字。G1 通过后，Baseline 的几何、网格、边界条件不得再改——若必须修改，需新开版本号并说明对实验对比的影响，原则上应回退到 G1 之前重新完成实验前准备。

Fatal 风险是 G1 未通过就开始实验，或实验后偷偷改 Baseline FE。无 Gate 记录会导致论文方法链不可审计。

必存数据：Baseline_FE_v0.wbpz。必存文件：Gate_G1.pdf（或同等门禁记录）。""",
            ),
        ],
    },
}

# Continue with flat phases B through I
STEPS.update({
    "阶段B：实验一次收全": {
        "B0 实验前": [
            (
                "B0-1",
                "SSS 核对、四向拍照与轴位截图",
                """正式实验开始前，按 SSS_Checklist 逐项核对机床状态，对机床进行四个方向的现场拍照，并截取数控系统轴位画面（含 X/Y/Z 坐标），确保与 FE 模型轴位一致。拖链、电缆、工装等细节必须与仿真一致；任何与 SSS 不符项应纠正或更新 SSS 并重新签字，不可带着不一致状态采集「正式数据」。

只拍一次轴位、或实验过程中擅自移动轴位而不记录，会使后续 C/D/E 阶段无法证明数据一致性。Fatal 风险是拖链/电缆姿态与 FE 不一致导致频率整体偏移。

必存数据：SSS_Checklist.pdf（签字版）。必存图片：四向现场照片 + 轴位截图（每次实验均需）。""",
            ),
            (
                "B0-2",
                "环境记录与静置",
                """记录实验室温度、湿度、主要干扰源（相邻机床是否运行等），机床在 SSS 状态下静置足够时间使温度分布稳定。环境记录写入 Environment_Log.csv，与每次实验编号关联。热变形与环境噪声会影响低频模态与 FRF 底噪。

常见遗漏是无环境 log。易错点是相邻设备仍在运行却未记录，导致某次 trial 噪底异常无法解释。

必存数据：Environment_Log.csv。本步骤不强制额外图片。""",
            ),
            (
                "B0-3",
                "背景噪底测试",
                """在不施加激励条件下采集至少 30 s 背景信号，确认后续激励测量的峰值高于背景至少 20 dB。若达不到，应排查干扰源、提高反光质量或调整激光设置，不可带着不合格信噪比开始正式扫频。

没测背景噪底则无法证明 FRF 峰值的可靠性。易错点是没有定量合格线，仅凭主观感觉「看起来有峰」。

必存数据：Background_Noise.csv。必存图片：背景噪底频谱图。""",
            ),
        ],
        "B1 步骤一：Y 中位 + SLDV 区 A（仅 P1–P9）": [
            (
                "B1-1",
                "SLDV 系统 Setup、试扫与 SNR 确认",
                """完成 SLDV 硬件与软件 Setup，保存 System_Config.txt（含软件版本、采样率、平均次数、跟踪模式等）。在代表性测点做试扫，确认 SNR 满足后续 Modal_ID_Settings 要求。距离过远或反光不良会导致数据整批报废。

常见遗漏是没有 Modal_ID_Settings 文件。易错点是试扫通过后正式测量却更改了平均次数或带宽设置而未记录。

必存数据：System_Config.txt。必存图片：Setup 现场照片。""",
            ),
            (
                "B1-2",
                "划定扫描区 A（polygon），不含 P10–P12",
                """在 SLDV 软件中划定扫描区 A 的多边形范围，该区用于 P1–P9 标定数据的场振型提取，必须不包含 P10–P12 盲验点位置。记录 polygon 顶点坐标与最小测距 d_min，存入 ScanZone_A.csv。若全场扫描包含盲验区，则盲验在逻辑上已被「看见」，失效。

Fatal 风险是扫描区 A 包含盲验点。易错点是没有保存 polygon 文件，事后无法证明区 A 范围。

必存数据：ScanZone_A.csv。必存图片：区 A 示意图（标注边界与盲验点位置）。""",
            ),
            (
                "B1-3",
            "3 种激励 × 5 次重复 Trial（激励点 EX1 / EX2 / EX3）",
            """在区 A 内采用 3 种激励配置，分别对应激励点 EX1、EX2、EX3（即 0-6 中定义的 E1–E3 激励位置；下文用 EX 前缀避免与阶段 E「盲验」步骤编号混淆）。每种配置做 5 次重复测量。每一次 trial 填写 Trial_Log_B1.csv（力幅、激励点 EX 编号、操作者、异常备注）。原始数据存入 02_Test_Raw/ExpB1_C0_SLDV_Raw/，文件夹结构在 Experiment_Matrix 中登记。

只做 3 次重复无法评估稳定性，论文也难以给出误差棒。易错点是各 trial 力幅差异过大，或三种激励实际打在相近位置导致某阶模态激励不足。""",
            ),
            (
                "B1-4",
                "力锤线性检验（轻敲 / 重敲）",
                """在代表性测点做轻敲与重敲对比，若峰值频率差异大于 2%，则正式测量仅采用轻敲范围内的力幅。非线性会导致假模态或频率漂移，影响配对与修正。

常见遗漏是完全不做线性检验。易错点是重敲导致结合部间隙开合，出现看似「清晰」但非线性的峰。

必存数据：Linearity_Check.pdf。必存图片：轻/重敲 FRF 对比图（可选）。""",
            ),
            (
                "B1-5",
                "EMA P1–P9；P1–P3 同点 SLDV 交叉验证",
                """对 P1–P9 进行传统 EMA（Experimental Modal Analysis，实验模态分析）测量；至少 P1–P3 与 SLDV 同点同步测量，用于交叉验证。EMA 数据作为 SCI 储备与硕士论文可信度补充，不可完全依赖单一 SLDV 手段。

完全不做 EMA 不会直接 Fatal，但会降低论文方法冗余度与毕业后 SCI 扩展空间。易错点是 EMA 与 SLDV 时间不同步、状态不一致。

必存数据：E3_C0_EMA/ 原始数据。本步骤图片可在 C 阶段统一整理。""",
            ),
            (
                "B1-6",
                "当晚双备份 raw 数据",
                """实验当日完成两份独立备份（例如本地工作站 + 网盘/移动硬盘），备份路径写入 Experiment_Matrix，原始文件夹设为只读。只存本机、或覆盖原始文件，一旦磁盘故障数据无法恢复，属于不可接受风险。

必存数据：备份路径记录（含校验信息如文件大小、哈希可选）。本步骤无需图片。""",
            ),
        ],
        "B2 步骤二：Y 极限（C1 raw）": [
            (
                "B2-1",
                "Y 极限轴位下重复 B1 流程（仅 raw，禁止 refit）",
                """实验当日必须先完成 B0-1~B0-3（SSS-H 核对、环境记录、背景噪底），不可引用 B1 当天记录。将机床移至 Y 极限轴位，按 SSS-H 更新拖链/电缆姿态并拍照存档，在相同 SLDV 设置下对区 A 进行测量。数据存入 02_Test_Raw/ExpB2_C1_SLDV_Raw/，填写 Trial_Log_B2.csv，实验当日完成双备份 Backup_Log_B2.txt。

本步骤仅采集 raw 数据供 F 阶段与 FE_Y_Extreme(θ*) 正推对比，严禁用 B2 数据调整 θ*。Fatal 风险是用 B2 数据 refit 后再声称「构型泛化验证」。易错点是仍用 SSS-M 检查表、或 Y 极限拖链姿态与 FE_Y_Extreme 不一致。""",
            ),
        ],
        "B3 步骤三：P10 / P11 / P12 独立盲测": [
            (
                "B3-1",
                "确认修正尚未使用盲验数据并封存",
                """实验当日必须先完成 B0-1~B0-3。在开始 P10–P12 测量前，确认 D 阶段修正尚未开始或至少尚未读取盲验文件夹。盲验 raw 分别存入 02_Test_Raw/Blind_P10/、Blind_P11/、Blind_P12/，在 θ* 冻结且 G3 通过之前保持只读封存。

盲验名存实亡（修正过程中偷看盲验数据）是 Fatal 级学术不端风险。""",
            ),
            (
                "B3-2",
                "三区各 5 次重复测量",
                """P10 位于结合部附近，P11 位于刀位旁（不安装传感器处），P12 位于空间远端；每点 5 次重复，激励方式与 B1 保持一致。填写 Trial_Log_B3.csv，实验当日完成双备份 Backup_Log_B3.txt。三点空间分散，不可挤在同一条梁面。

仍只测 1 个盲验点无法支撑「盲验通过」结论。""",
            ),
            (
                "B3-3",
                "盲验封存 / 开封记录",
                """测量完成后对盲验文件夹封存，填写 Seal_Record.pdf（封存日期、见证人）。开封记录 Unseal_Record.pdf 仅在 θ* 冻结且 G3 通过后、阶段 E 步骤 E1 中填写，B3 阶段不得填写开封记录。

无见证签字或开封早于冻结，答辩时会被质疑程序无效。""",
            ),
        ],
        "B4 备用重测": [
            (
                "B4-1",
                "仅重测失败步骤（B4 备用方案）",
                """若 B1、B2、B3 任一步数据不合格（SNR 不足、重复性超标、操作失误等），启用 B4 备用时段仅重测失败步骤，并在 Experiment_Matrix 中新建实验编号，保留失败数据不得覆盖。严禁用修正后的模型回头「指导」重测实验参数以凑合格结果。

没预留 B4 会导致时间压力下接受劣质数据。Fatal 风险是用 Updated FE 回头调整实验设置却声称是 Baseline 实验流程。

本步骤无固定必存文件，但重测记录必须完整。""",
            ),
        ],
    },
    "阶段C：实验模态识别与配对": [
        (
            "C1",
            "稳态图识别并筛选 6 阶弹性模态",
            """对 B1 的 SLDV/EMA 数据做模态识别，采用与 Modal_ID_Settings 一致的带宽、模型阶次与稳态图判据，提取稳定极点并筛选 6 阶弹性模态。剔除的极点应记录原因（不稳定、重复、物理不合理等），输出 Modal_Parameters.csv（频率、阻尼、参与因子等）。

假峰当真模态会连锁影响配对与修正。易错点是不同 trial 识别阶次不一致却强行取平均——应先用统一设置识别再比较。

必存数据：Modal_Parameters.csv。必存图片：稳态图（Stabilization Diagram）。""",
        ),
        (
            "C2",
            "Exp ↔ FE 配对（频率 + 法向 MAC + 振型）",
            """在 C1 得到的 6 阶实验模态与 A4-3 的 6 阶 FE 模态之间，综合频率接近程度、法向 MAC 与振型物理形态做配对，输出 Mode_Pairing_Table.xlsx。存在配对歧义时（两阶频率接近），必须记录取舍理由，不可默认按阶号 1 对 1、2 对 2。

按阶号硬对是常见 Fatal 易错点。错配后越修越偏，MAC 图看起来「更好」其实是错配换阶。

必存数据：Mode_Pairing_Table.xlsx。必存图片：实验 Mode 1–6 振型图。""",
        ),
        (
            "C3",
            "5 次重复性统计（mean ± std）",
            """对 5 次 trial 的频率与 MAC 做 mean ± std 统计，论文主结果报告 Trial 1–3，但 Repeatability.csv 应保留全部 5 次信息供 SCI 扩展。重复性标准：频率 std 小于 1.5%（优秀小于 1%）。

只报一次测量值无误差棒，结论可信度不足。易错点是剔除异常 trial 却未按 Trial_Selection_Rule 执行。

必存数据：Repeatability.csv。必存图片：重复性误差棒图。""",
        ),
        (
            "C4",
            "G2 配对表冻结签字",
            """配对表经导师或第二人复核后冻结，填写 Gate_G2.pdf。G2 通过后不得随意更改配对关系；若必须更改，需说明原因并评估对 D/E 阶段的影响，原则上应重新走 C 阶段。

一人配对无复核、无签字，是方法链薄弱环节。必存文件：Gate_G2.pdf。""",
        ),
    ],
    "阶段D：模型修正（仅用 P1–P9）": [
        (
            "D1",
            "修正前初对比 Before（频率 + MAC + FRF）",
            """在 Baseline FE 上与实验对比，输出 Freq_MAC_Before.csv 与 Before 对比图，定量展示修正必要性。若配对本身存疑，应先回 C2 解决，不可在错配基础上开始优化。

跳过 Before 对比直接修，论文缺少「问题陈述」依据。易错点是配对错仍强行修正。

必存数据：Freq_MAC_Before.csv。必存图片：Before 对比图（频率 + MAC）。""",
        ),
        (
            "D2",
            "建立目标函数 J(θ)，仅用 P1–P9",
            """目标函数 J(θ) 仅纳入 P1–P9 标定点的频率与法向 MAC 误差，权重设置写入 Objective_Function.txt。严禁将 P10–P12 或 SLDV 全场扫描数据纳入 J(θ)——否则盲验失效，属于 Fatal 级错误。SLDV 场数据可用于可视化，不可用于驱动参数优化。

用全场 MAC 优化是常见 Fatal 易错点。必存文件：Objective_Function.txt。""",
        ),
        (
            "D3",
            "从 4 参数试修，再扩展至 6–8 个",
            """按灵敏度排序从 4 参数开始试修，记录 Iteration_Log.csv，观察 J(θ) 收敛曲线；用 L-curve 等方法选择有效参数个数，再扩展至 6–8 个，全程参数必须在 Bounds 内。一次同时优化 8 个以上参数易导致盲验崩溃。

无迭代 log 无法证明收敛过程可信。必存数据：Iteration_Log.csv。必存图片：J(θ) 收敛曲线。""",
        ),
        (
            "D4",
            "2–3 组初值 + SCI 储备 10-start",
            """至少用 2–3 组不同初值验证 θ* 非唯一性风险；同时完成 10-start 全局搜索并存档于 10start/，供毕业后 SCI 讨论解的唯一性。只做 1 组初值可能导致 θ* 落在局部最优。

初值导致负刚度或非物理参数时应排查 Bounds 与参数化方式。必存数据：10start/ 文件夹。""",
        ),
        (
            "D5",
            "收敛 θ* 并通过 G3 冻结",
            """θ* 收敛后导出 Updated_FE_ThetaStar.wbpz，填写 Gate_G3.pdf。G3 通过后 θ* 与 Updated FE 冻结，不得再微调参数；谐响应与盲验均基于该版本。修正过程中不得修改 Baseline 几何或网格来「配合」实验。

G3 后仍微调 θ 是 Fatal 风险。必存数据：Updated_FE_ThetaStar.wbpz。必存图片：修正前后 MAC 对比图。""",
        ),
    ],
    "阶段E：盲验开封与判定": [
        (
            "E1",
            "填写开封记录，确认 θ* 已冻结",
            """在 G3 通过且 θ* 确认冻结后，填写 Unseal_Record.pdf，见证人签字，方可打开 Blind_P10/11/12/ 原始数据。开封日期必须晚于 θ* 冻结日期，程序链完整可查。

开封后仍改 θ 使盲验失去意义。必存文件：Unseal_Record.pdf。""",
        ),
        (
            "E2",
            "P10 / P11 / P12 三点 Updated / Baseline / Exp 对比",
            """对每一点分别计算 Updated(θ*)、Baseline 与实验 Exp 的频率与法向 MAC，输出 Blind_Summary.csv。必须三点都报，不可只报 P10 却说「盲验通过」。

单点偶然合格不能代表模型泛化能力。必存图片：三点盲验汇总图。""",
        ),
        (
            "E3",
            "汇总判定：≥2/3 合格（优秀 3/3），含频率与 MAC",
            """盲验单点判定需同时报告频率误差与法向 MAC（两点均记录于 Blind_Summary_3points.csv）。MAC 合格线：≥0.55（优秀 ≥0.60）；频率误差合格线：≤10%（优秀 ≤8%）。单点判定为 MAC 与频率均达标；三点中至少 2 点达标为通过，优秀为 3/3 全通过。报告平均 MAC ± std 与平均频率误差。

1 点 pass 即下结论是常见错误。Fail 时不得用盲验数据回头调参。""",
        ),
        (
            "E4",
            "G4 签字；fail 则回 D3 且禁用盲验数据",
            """填写 Gate_G4.pdf。若盲验 fail，只能回 D3 用 P1–P9 重新修正，严禁使用 P10–P12 数据参与任何优化。Fail 分析与尝试记录必须保留，不可隐瞒 fail 仍写「盲验成功」。

必存文件：Gate_G4.pdf；若有 fail，另存 Fail_Analysis.pdf。""",
        ),
    ],
    "阶段F：C1 构型正推验证": [
        (
            "F1",
            "FE_Y_Extreme(θ*) 正推（禁止 refit）",
            """使用冻结的 θ* 在 FE_Y_Extreme 模型上正推模态结果，输出 C1_Forward.csv。禁止用 B2 raw 数据对 θ* 做任何 refit——B2 仅用于对比验证。

用中位 FE 对比 Y 极限实验是逻辑 Fatal 错误。必存数据：C1_Forward.csv。""",
        ),
        (
            "F2",
            "与 B2 raw 对比频率 / MAC",
            """将 F1 正推结果与 B2 实验 raw 对比。若误差较大，在论文 Limitation 中如实说明，降级表述为「趋势一致」或「部分模态吻合」，不可误差很大仍声称「构型泛化验证成功」。

必存：对比结论写入论文第 4–5 章 Limitation；数据对比表可并入 C1_Forward.csv。""",
        ),
    ],
    "阶段G：刀位谐响应仿真与实验": {
        "G1 谐响应仿真": [
            (
                "G1-1",
                "Updated + Baseline 谐响应，同激励频段",
                """在 Updated(θ*) 与 Baseline FE 上分别做谐响应（Harmonic Response）分析，激励频段与幅值与 G2 实验一致，必须基于 θ* 版本。不可用 Baseline 结果冒充 Updated 的谐响应改进。

必存数据：Harmonic_Updated.csv、Harmonic_Baseline.csv。""",
            ),
            (
                "G1-2",
                "TP 输出 FRF；模态叠加收敛性",
                """在 TP（刀位点）输出 FRF，检查模态叠加阶次收敛，记录使用了多少阶模态，输出 Mode_Convergence.csv。阶次不足会导致 FRF 缺峰。

必存数据：Mode_Convergence.csv。""",
            ),
        ],
        "G2 谐响应实验": [
            (
                "G2-1",
                "TP 安装加速计并记录附加质量",
                """在 TP 安装加速计，精确记录附加质量与安装位置；仿真中施加等效质量或在 Limitation 说明未建模影响。附加质量会改变局部高频特性。

必存文件：Added_Mass_Note.pdf。必存图片：传感器布置照片。""",
            ),
            (
                "G2-2",
                "测输入力 + TP / 参考点 FRF（激励方式定稿）",
                """谐响应实验采用与 G1 仿真一致的频段。激励推荐使用激振器+力传感器或力锤（二选一，写入 Harmonic_Excitation_Spec.pdf），激励位置优先 EX1 或 TP 附近（与仿真载荷施加方式对应）。测量输入力谱与 TP、参考点 FRF，力链标定结果存档 Force_Chain.pdf。只测加速度不测力，FRF 幅值不可信。""",
            ),
            (
                "G2-3",
                "5 重复；相干系数 γ² ≥ 0.7",
                """谐响应实验 5 次重复，论文报告 Trial 1–3；每次计算相干系数 γ²（Coherence），合格线 γ² ≥ 0.7。γ² 低的频段不应拿来做峰值对比。

必存数据：TP_FRF_5trials.csv。""",
            ),
            (
                "G2-4",
                "G5 签字",
                """谐响应主指标为峰值频率对比（Updated vs Baseline vs Exp），合格 ≤12%（优秀 ≤10%）。填写 Gate_G5.pdf，Limitation 中说明谐响应不等同于切削稳定性实验。

必存文件：Gate_G5.pdf。必存图片：谐响应三线 FRF 对比图（Updated / Baseline / Exp）。""",
            ),
        ],
    },
    "阶段H：论文撰写": [
        (
            "H1",
            "第 1、2 章（绪论、理论基础）",
            """先完成第 1、2 章，定稿 3 条创新点，明确研究边界与 Limitation。方法章不能缺失；创新点应可追溯到 G1–G5 门禁与盲验设计。

必存：Ch1-2.docx。必存图片：技术路线图。""",
        ),
        (
            "H2",
            "第 3 章（随仿真/实验补图）",
            """第 3 章随进展补充图表，维护 Figure_Index.xlsx。核心图不少于 16 张，建议初稿至少包含：技术路线、CAD/FE 模型、网格收敛、重力变形、FE 前 6 阶振型、测点布置、稳态图、实验振型、Before/After MAC、J(θ) 收敛、盲验汇总、构型正推（若做 F）、谐响应三线 FRF、重复性误差棒、灵敏度龙卷风、SSS 状态图。图号、数据路径、生成脚本应可追溯。""",
        ),
        (
            "H3",
            "第 4–6 章（方法、结果、结论）+ Limitation 清单",
            """第 4–6 章随 C/D/E/G 阶段填充，结论必须与门禁结果一致。附录 A–D（坐标变换、MAC 协议、Trial 规则、盲验程序）应同步完成。Limitation 必写清单（写入 Limitation_Checklist.pdf 并体现在第 5–6 章）：① TP 附加质量影响；② 谐响应≠切削稳定性；③ 单机床/单构型外推限度；④ 构型正推（F 阶段）误差范围；⑤ 结合部参数非唯一性（多初值）；⑥ SSS 无法完全复现的项。""",
        ),
        (
            "H4",
            "符号表与参考文献",
            """符号表 Symbols.docx 与正文一致；参考文献 GB/T 7714 格式。符号多义是格式打回常见原因。""",
        ),
        (
            "H5",
            "涉密脱敏 Public 版",
            """若企业 CAD 涉密，制作 Public 版 FE 与插图，全文可送审。未做 Public 版可能导致不能送审。""",
        ),
        (
            "H6",
            "与 SCI 章节划分规划",
            """规划硕士论文与毕业后 SCI 论文的章节划分，避免自我抄袭查重问题。SCI 可侧重盲验、多初值、EMA 交叉验证等扩展内容。""",
        ),
    ],
    "阶段I：SCI 数据储备（硕士可少写、数据必须先收）": [
        (
            "I1",
            "B1 / B2 / 10-start 算完存档",
            """B1、B2 原始数据与 10-start 优化结果在 θ* 冻结后立即存档于 04_SCI_Reserve/，与硕士论文主结论隔离。毕业后写 SCI 时不应需要重跑 ANSYS 或重测机床。

易错点：B2 数据覆盖或混淆 θ* 版本。""",
        ),
        (
            "I2",
            "5 重复 raw 全保留",
            """02_Test_Raw/ 保留全部 5 次 trial raw，论文虽只用 Trial 1–3，但 SCI 可能需要置信区间分析。只存 processed 删 trial 4–5 不可逆。

必存：02_Test_Raw/ 完整备份。""",
        ),
        (
            "I3",
            "EMA、C1、三盲验汇总（SCI 扩展）",
            """EMA 交叉验证、C1 构型正推、三盲验详细汇总为 SCI 储备，硕士论文可精简表述，但原始数据必须在 B/C/E/F 阶段已采集完毕。不写 SCI 不等于可以不测。""",
        ),
    ],
})


def add_criteria_section(doc: Document) -> None:
    doc.add_paragraph("附录：合格指标与测点说明", style="Heading 1")
    add_para(
        doc,
        """下列指标为全文统一的合格线，与门禁 G1–G5 配合使用。Simulation（仿真）侧在 A 阶段核对，Experiment（实验）侧在 C/E/G 阶段核对。表述中「优秀」为更高标准，硕士论文至少应满足「合格」列。""",
    )
    criteria = [
        "质量误差（FE 总质量 vs BOM/称重）：合格小于 5%，优秀小于 3%。",
        "网格收敛（前三阶频率相对变化 Δf₁₋₃）：合格小于 2%。",
        "实验重复性（5 次 trial 频率标准差）：合格小于 1.5%，优秀小于 1%。",
        "标定点 P1–P9 频率误差（Updated vs Exp）：合格 ≤8%，优秀 ≤5%。",
        "标定点 P1–P9 MAC（法向投影后）：合格 ≥0.65，优秀 ≥0.70。",
        "盲验单点 MAC（P10/P11/P12 各点）：合格 ≥0.55，优秀 ≥0.60。",
        "盲验单点频率误差（Updated vs Exp，各点）：合格 ≤10%，优秀 ≤8%。",
        "盲验通过率：合格为三点中 MAC 与频率均达标的点数 ≥2/3，优秀为 3/3。",
        "谐响应峰值频率误差（Updated vs Exp）：合格 ≤12%，优秀 ≤10%。",
    ]
    for c in criteria:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_paragraph("测点与激励点分工", style="Heading 2")
    add_para(
        doc,
        """P1–P9 为标定点，参与模型修正目标函数 J(θ)，应在结构上分散布置并优先选择对结合部参数敏感的位置。P10、P11、P12 为盲验点，不参与任何修正优化；P10 建议位于结合部附近，P11 位于刀位旁但不在传感器安装位置，P12 位于空间远端以检验远距离预测能力。三者必须在 Pre-registration 中事先登记并在空间上分散，不可位于同一梁段。

E1–E3（文档中亦记 EX1–EX3）为激励点，应在结构上分散。TP 为刀位点，用于谐响应输出对比，映射与附加质量记录必须完整。""",
    )


def add_fatal_section(doc: Document) -> None:
    doc.add_paragraph("附录：Fatal 级致命遗漏清单", style="Heading 1")
    add_para(doc, "下列任一条发生，将直接威胁盲验成立性、数据可追溯性或答辩可信度。实验/仿真过程中逐条自查：")
    fatals = [
        "G1 未通过即开始正式实验，或实验后擅自修改 Baseline 几何/网格/BC。",
        "P10–P12 或 SLDV 全场数据进入修正目标函数 J(θ)（修正仅用 P1–P9）。",
        "θ* 冻结（G3）之前打开、查看或分析盲验 raw 数据。",
        "B1 扫描区 A 的 polygon 包含 P10–P12 盲验点位置。",
        "用 B2（Y 极限）raw 数据 refit θ*，再声称构型泛化验证。",
        "实验模态与 FE 模态按阶号 1 对 1 硬配对，未做频率+MAC+形态综合配对。",
        "结合部全刚性连接，或未做网格收敛即冻结 Baseline。",
        "用普通模态（无预应力）代替重力预应力模态与实验对比。",
        "测点映射距离 d>10 mm 未复核仍用于 MAC 与修正。",
        "G3 通过后仍微调 θ*，或盲验 fail 后使用盲验数据调参。",
        "盲验仅测 1 点，或无 Pre-registration 签字即开始 B3。",
        "无 Trial 选用规则，事后从 5 次重复中挑选「最好 3 次」。",
        "只存 3 次 trial processed 而删除 raw trial 4–5。",
        "谐响应结果未经说明即声称加工稳定性或切削性能结论。",
        "B2/B3 实验日未重走 B0（SSS/环境/噪底），却与 B1 数据混为同一状态。",
    ]
    for f in fatals:
        doc.add_paragraph(f, style="List Number")


def add_code_legend(doc: Document) -> None:
    doc.add_paragraph("附录：编号对照说明", style="Heading 1")
    add_para(
        doc,
        """为避免混淆，全文采用以下编号约定（阅读步骤码时请先对照本表）：

门禁 G1–G5：Baseline 冻结、配对冻结、θ* 冻结、盲验判定、谐响应判定五道签字门禁，与阶段 G「谐响应」无关。  
阶段 G 步骤 G1-1/G2-1 等：G 表示 Harmonic Response（谐响应）阶段内的步骤序号，不等于门禁 G1/G2。  
阶段 E 步骤 E1–E4：E 表示 Blind Validation（盲验判定）阶段内的步骤序号，不等于激励点 E1–E3。  
激励点 EX1–EX3：即坐标文件中的 E1–E3 激励位置；文档正文优先写 EX 前缀。  
实验文件夹 ExpB1/ExpB2：ExpB1= B1 实验 batch，ExpB2= B2 实验 batch，不等于阶段步骤编号。  
构型 C1：专指 Y 极限构型验证（F 阶段）；阶段 C 步骤 C1 指模态识别，二者含义不同。""",
    )


DELIVERABLES: dict[str, dict[str, list[str]]] = {
    "0-1": {
        "data": [
            "Experiment_Matrix.xlsx — 实验矩阵总表",
            "01_Admin/ — 管理类（矩阵、预约、安全、Pre-registration、门禁 PDF）",
            "02_Test_Raw/ — 全部实验 raw（只读）",
            "03_Simulation/ — Baseline、Updated、FE_Y_Extreme、谐响应仿真",
            "04_SCI_Reserve/ — 10-start、EMA、5 trial 完整备份",
            "05_Thesis/ — 论文、Figure_Index、附录、Public 版",
            "Folder_Structure.pdf — 目录树说明与命名规则",
        ],
        "images": [],
    },
    "0-2": {
        "data": ["P10-P12_PreRegistration.pdf — 盲验三点坐标、类型、实物标记方式、导师签字页"],
        "images": ["P10-P12_Marker_Photo.jpg — 盲验点实物标记照（可见编号与贴点位置）"],
    },
    "0-3": {
        "data": [
            "SSS-M_Definition.pdf — Y 中位构型标准状态（B1 + Baseline FE）",
            "SSS-H_Definition.pdf — Y 极限构型标准状态（B2 + FE_Y_Extreme）",
            "SSS-M_Checklist.pdf / SSS-H_Checklist.pdf — 各构型实验前核对表",
        ],
        "images": [
            "SSS-M_Overview.png — Y 中位状态示意",
            "SSS-H_Overview.png — Y 极限状态示意（拖链/电缆姿态）",
        ],
    },
    "0-4": {
        "data": ["MAC_Calculation_Protocol.pdf — 法向投影公式、归一化、阈值、版本号"],
        "images": ["MAC_Normal_Projection.png — 法向投影示意（可选但建议存）"],
    },
    "0-5": {
        "data": ["Trial_Selection_Rule.pdf — Trial 1–3 选用与异常 trial 剔除规则"],
        "images": [],
    },
    "0-6": {
        "data": ["Coordinate_Transform.pdf — 机床/CAD/FE/SLDV 四套坐标变换关系与验证记录"],
        "images": ["Coordinate_Systems.png — 坐标系关系示意图（原点、轴向）"],
    },
    "0-7": {
        "data": ["Mode_Screening_Rules.docx — 12 阶提取、6 阶弹性模态筛选与剔除判据"],
        "images": ["Mode_Screening_Example.png — 被剔除模态 vs 保留模态振型对比示例"],
    },
    "0-8": {
        "data": ["Booking_Record.pdf — 机床/SLDV/B4 备用时段预约记录（邮件或系统截图）"],
        "images": [],
    },
    "0-9": {
        "data": [
            "Purchase_List.pdf — 反光膜、加速计、力传感器等采购清单",
            "Laser_Safety_Signoff.pdf — 激光安全告知签字页",
        ],
        "images": [],
    },
    "A1-1": {
        "data": [
            "A1_Model.step — CAD/几何源文件",
            "Simplification_List.docx — 几何简化清单（删/留部件及理由）",
        ],
        "images": ["CAD_Assembly.png — CAD 装配图（标注主要部件与坐标原点）"],
    },
    "A1-2": {
        "data": ["Mass_Summary.csv — FE 质量/质心与 BOM 或称重对比表"],
        "images": ["Mass_CoG_Compare.png — 质心位置对比示意图"],
    },
    "A2-1": {
        "data": ["Joint_Theta0.csv — 结合部参数初值 θ₀、Bounds、物理含义与来源"],
        "images": ["Joint_Locations.png — 结合部位置示意图（标注参数对应位置）"],
    },
    "A2-2": {
        "data": ["Boundary_Spec.pdf — 地脚 BC 类型、位置、与 SSS 对应说明"],
        "images": ["BC_Annotation.png — FE 模型边界条件标注图"],
    },
    "A3-1": {
        "data": ["Mesh_Convergence.csv — 粗/中/细网格前三阶频率对比表"],
        "images": [
            "Mesh_Views.png — 网格划分图（整体 + 结合部局部）",
            "Mesh_Convergence_Curve.png — 前三阶频率收敛曲线",
        ],
    },
    "A4-1": {
        "data": ["Static_Results.rst — 重力静力分析结果文件"],
        "images": ["Static_Deformation.png — 重力变形云图"],
    },
    "A4-2": {
        "data": [
            "Modal_Results.rst — 预应力模态分析结果",
            "Modal_Settings.txt — 模态分析设置（阶次、频率范围、预应力链接）",
        ],
        "images": [],
    },
    "A4-3": {
        "data": ["Elastic_Modes_6.csv — 6 阶弹性模态频率及剔除记录"],
        "images": ["FE_Mode1-6.png — FE 前 6 阶弹性模态振型图（可 6 张或拼图）"],
    },
    "A5-1": {
        "data": ["P1-P12_Mapping.csv — 测点/激励点/TP 映射节点与距离 d"],
        "images": ["Point_Layout.png — 测点布置图（FE 与实验对照）"],
    },
    "A5-2": {
        "data": ["Sensitivity.csv — 参数灵敏度分析结果"],
        "images": ["Sensitivity_Tornado.png — 灵敏度龙卷风图"],
    },
    "A6-1": {
        "data": ["FE_Y_Extreme.wbpz — Y 极限构型 FE 工程文件"],
        "images": ["Y_Extreme_Config.png — Y 极限轴位与拖链/电缆姿态截图或示意图"],
    },
    "A7-1": {
        "data": [
            "Baseline_FE_v0.wbpz — 冻结 Baseline 有限元工程",
            "Gate_G1.pdf — G1 门禁签字记录",
        ],
        "images": [],
    },
    "B0-1": {
        "data": ["SSS_Checklist_Signed.pdf — 当次实验 SSS 核对签字版"],
        "images": [
            "Machine_4Views.jpg — 机床四向现场照片",
            "AxisPosition_Screenshot.png — 数控系统轴位截图（含 X/Y/Z）",
        ],
    },
    "B0-2": {
        "data": ["Environment_Log.csv — 温度、湿度、干扰源、静置时间记录"],
        "images": [],
    },
    "B0-3": {
        "data": ["Background_Noise.csv — 30 s 背景噪底数据"],
        "images": ["Background_Spectrum.png — 背景噪底频谱图"],
    },
    "B1-1": {
        "data": [
            "System_Config.txt — SLDV 软硬件版本、采样率、平均次数等 Setup 参数",
            "Modal_ID_Settings.txt — 模态识别软件设置（与 C1 一致）",
        ],
        "images": ["SLDV_Setup_Photo.jpg — SLDV Setup 现场照片"],
    },
    "B1-2": {
        "data": ["ScanZone_A.csv — 区 A polygon 顶点坐标与 d_min"],
        "images": ["ScanZone_A_Map.png — 区 A 示意图（边界 + 盲验点位置）"],
    },
    "B1-3": {
        "data": [
            "02_Test_Raw/ExpB1_C0_SLDV_Raw/ — B1 全部 raw（3 激励 EX1–EX3 ×5 trial，只读）",
            "Trial_Log_B1.csv — 每次 trial 力幅、EX 编号、操作者、异常备注",
        ],
        "images": [],
    },
    "B1-4": {
        "data": ["Linearity_Check.pdf — 力锤轻/重敲线性检验记录与结论"],
        "images": ["Linearity_FRF_Compare.png — 轻/重敲 FRF 对比（建议存）"],
    },
    "B1-5": {
        "data": [
            "E3_C0_EMA/ — EMA P1–P9 原始数据",
            "CrossCheck_P1-P3.csv — P1–P3 同点 SLDV/EMA 交叉验证记录",
        ],
        "images": [],
    },
    "B1-6": {
        "data": [
            "Backup_Log_B1.txt — 双备份路径、日期、文件大小/校验信息",
            "Experiment_Matrix 更新 — 登记 B1 raw 与备份位置",
        ],
        "images": [],
    },
    "B2-1": {
        "data": [
            "02_Test_Raw/ExpB2_C1_SLDV_Raw/ — Y 极限 SLDV raw（禁止 refit θ*）",
            "Trial_Log_B2.csv — B2 各 trial 记录",
            "Backup_Log_B2.txt — 当日双备份路径与校验",
            "B0_Log_B2.pdf — 当日 B0-1~B0-3 记录（SSS-H、环境、噪底）",
        ],
        "images": [
            "Y_Extreme_Axis.png — Y 极限轴位截图",
            "Y_Extreme_4Views.jpg — 四向现场照片（SSS-H）",
        ],
    },
    "B3-1": {
        "data": [
            "Blind_P10/ — P10 盲验 raw（封存）",
            "Blind_P11/ — P11 盲验 raw（封存）",
            "Blind_P12/ — P12 盲验 raw（封存）",
        ],
        "images": ["BlindPoints_Layout.jpg — P10–P12 位置实物图（与 Pre-registration 对照）"],
    },
    "B3-2": {
        "data": [
            "02_Test_Raw/Blind_P10/Trial1-5/ — P10 五次重复 raw",
            "02_Test_Raw/Blind_P11/Trial1-5/ — P11 五次重复 raw",
            "02_Test_Raw/Blind_P12/Trial1-5/ — P12 五次重复 raw",
            "Trial_Log_B3.csv — 盲验各 trial 记录",
            "Backup_Log_B3.txt — 当日双备份路径与校验",
        ],
        "images": [],
    },
    "B3-3": {
        "data": ["Seal_Record.pdf — 盲验封存记录（日期、见证人；开封记录在 E1 填写）"],
        "images": [],
    },
    "B4-1": {
        "data": [
            "B4_Retest_Log.pdf — 重测原因、新实验编号、与失败 batch 关系",
            "（若状态变化）补 B0_Log_B4.pdf — 当日 SSS/环境/噪底",
        ],
        "images": [
            "B4_Retest_Photos.jpg — 仅当重测因机床状态变化时：四向照 + 轴位截图",
        ],
    },
    "C1": {
        "data": [
            "Modal_Parameters.csv — 实验识别模态参数（频率、阻尼等）",
            "Pole_Rejection_Log.csv — 剔除极点清单与原因",
        ],
        "images": ["Stabilization_Diagram.png — 稳态图"],
    },
    "C2": {
        "data": ["Mode_Pairing_Table.xlsx — Exp↔FE 配对表（含歧义说明）"],
        "images": ["Exp_Mode1-6.png — 实验前 6 阶模态振型图"],
    },
    "C3": {
        "data": ["Repeatability.csv — 5 次 trial 频率/MAC 的 mean±std"],
        "images": ["Repeatability_Bars.png — 重复性误差棒图"],
    },
    "C4": {
        "data": ["Gate_G2.pdf — G2 配对表冻结签字记录"],
        "images": [],
    },
    "D1": {
        "data": ["Freq_MAC_Before.csv — 修正前 Baseline 与实验频率/MAC 对比"],
        "images": ["Before_Compare.png — Before 频率 + MAC 对比图"],
    },
    "D2": {
        "data": ["Objective_Function.txt — J(θ) 定义、权重、仅用 P1–P9 的声明"],
        "images": [],
    },
    "D3": {
        "data": [
            "Iteration_Log.csv — 每轮迭代参数、目标函数值、是否满足 Bounds",
            "Lcurve_Selection.pdf — L-curve 选参记录",
        ],
        "images": ["J_Convergence.png — J(θ) 收敛曲线"],
    },
    "D4": {
        "data": [
            "10start/ — 10 组初值优化结果（SCI 储备）",
            "MultiStart_Summary.csv — 多初值 θ* 对比汇总",
        ],
        "images": ["MultiStart_Boxplot.png — 多初值结果 boxplot（SCI 储备，建议存）"],
    },
    "D5": {
        "data": [
            "Updated_FE_ThetaStar.wbpz — 修正后冻结 FE 工程",
            "Gate_G3.pdf — G3 门禁签字记录",
            "ThetaStar_Frozen.csv — θ* 参数终值与 Bounds",
        ],
        "images": ["MAC_Before_After.png — 修正前后 MAC 对比图"],
    },
    "E1": {
        "data": ["Unseal_Record.pdf — 盲验开封记录（见证签字，日期晚于 G3）"],
        "images": [],
    },
    "E2": {
        "data": ["Blind_Summary.csv — P10/P11/P12 各点 Updated/Baseline/Exp 对比"],
        "images": ["Blind_3Points_Summary.png — 三点盲验汇总图"],
    },
    "E3": {
        "data": ["Blind_Summary_3points.csv — 通过率判定表（≥2/3 或 3/3）"],
        "images": ["Blind_Representative_FRF.png — 代表点 FRF 对比图"],
    },
    "E4": {
        "data": [
            "Gate_G4.pdf — G4 门禁签字记录",
            "Fail_Analysis.pdf — 若 fail：原因分析与回 D3 记录（禁用盲验数据）",
        ],
        "images": [],
    },
    "F1": {
        "data": ["C1_Forward.csv — FE_Y_Extreme(θ*) 正推模态结果"],
        "images": [],
    },
    "F2": {
        "data": ["C1_Compare_Summary.csv — 正推 vs B2 raw 频率/MAC 对比表"],
        "images": ["C1_Compare_Chart.png — 构型验证对比图（误差大时仍须存）"],
    },
    "G1-1": {
        "data": [
            "Harmonic_Updated.csv — Updated(θ*) 谐响应 TP 输出",
            "Harmonic_Baseline.csv — Baseline 谐响应 TP 输出",
        ],
        "images": [],
    },
    "G1-2": {
        "data": ["Mode_Convergence.csv — 模态叠加阶次收敛记录"],
        "images": ["Mode_Superposition_Conv.png — 叠加收敛曲线（建议存）"],
    },
    "G2-1": {
        "data": ["Added_Mass_Note.pdf — TP 加速计附加质量、安装位置、仿真等效说明"],
        "images": ["TP_Sensor_Photo.jpg — TP 传感器布置照片"],
    },
    "G2-2": {
        "data": [
            "Harmonic_Excitation_Spec.pdf — 谐响应激励方式（激振器/力锤、位置、频段）",
            "Force_Chain.pdf — 力链标定记录",
            "Input_Force_Spectrum.csv — 输入力谱数据",
        ],
        "images": [],
    },
    "G2-3": {
        "data": ["TP_FRF_5trials.csv — TP 谐响应 5 次重复 FRF 及 γ²"],
        "images": ["Coherence_Gamma2.png — 相干系数 γ² 图（建议存）"],
    },
    "G2-4": {
        "data": ["Gate_G5.pdf — G5 门禁签字记录"],
        "images": ["Harmonic_FRF_3Lines.png — Updated/Baseline/Exp 三线 FRF 对比图"],
    },
    "H1": {
        "data": ["Ch1-2.docx — 第 1、2 章草稿（含创新点 3 条）"],
        "images": ["Tech_Roadmap.png — 技术路线图"],
    },
    "H2": {
        "data": [
            "Figure_Index.xlsx — 图索引（图号、文件名、数据路径、脚本）",
            "Core_Figure_List.pdf — 16+ 核心图目录初稿（见 H2 步骤正文列表）",
        ],
        "images": ["Core_Figures/ — 按 Figure_Index 存放的全部核心论文图"],
    },
    "H3": {
        "data": [
            "Ch4-6.docx — 第 4–6 章",
            "Appendix_A-D.pdf — 坐标/MAC/Trial/盲验程序附录",
            "Limitation_Checklist.pdf — Limitation 必写清单（6 条）",
        ],
        "images": [],
    },
    "H4": {
        "data": ["Symbols.docx — 符号表", "References.bib — GB/T 7714 参考文献"],
        "images": [],
    },
    "H5": {
        "data": ["FE_Public/ — 脱敏 Public 版 FE 与可送审插图"],
        "images": ["Public_Figures/ — 涉密脱敏后的论文插图"],
    },
    "H6": {
        "data": ["SCI_Outline.docx — 硕士论文与 SCI 章节划分规划"],
        "images": [],
    },
    "I1": {
        "data": [
            "04_SCI_Reserve/B1_raw/ — B1 隔离备份",
            "04_SCI_Reserve/B2_raw/ — B2 隔离备份",
            "04_SCI_Reserve/10start/ — 多初值优化备份",
        ],
        "images": [],
    },
    "I2": {
        "data": ["02_Test_Raw/ — 全部 5 trial 原始数据（不可删 trial 4–5）"],
        "images": [],
    },
    "I3": {
        "data": [
            "SCI_EMA_Summary.csv — EMA 交叉验证汇总",
            "SCI_C1_Summary.csv — C1 构型验证汇总",
            "SCI_Blind_Detail.csv — 三盲验详细汇总",
        ],
        "images": [],
    },
}


CHECKPOINTS: dict[str, list[str]] = {
    "0-1": ["Experiment_Matrix 中每个实验编号唯一且已指定备份路径", "文件夹命名含阶段代号而非纯日期", "矩阵字段含操作者与只读封存标记"],
    "0-2": ["Pre-registration 日期早于 B1 正式实验", "P10/P11/P12 坐标与实物标记一致", "导师签字页已扫描存档"],
    "0-3": ["SSS 含轴位、伺服状态、拖链/电缆、辅机", "检查表可在 B0 快速复用", "与 FE 轴位截图可一一对应"],
    "0-4": ["法向投影公式与归一化方式已写清", "与第 3、4 章公式一致", "版本号已定稿"],
    "0-5": ["Trial 1–3 选取规则含剔除条件", "剔除需有 Trial_Log 依据", "规则已写入第 4 章方法"],
    "0-6": ["四套坐标系变换关系已验证", "单位统一 mm", "P1 试映射距离 d 已复核"],
    "0-7": ["12 阶提取与 6 阶筛选判据明确", "剔除模态有记录", "示例振型图已保存"],
    "0-8": ["B1–B3 时段已连续预约", "B4 备用时段已预留", "预约记录已存档"],
    "0-9": ["反光膜/传感器等到货", "激光安全告知已签字", "采购清单与实验需求一致"],
    "A1-1": ["FE 坐标系与 Coordinate_Transform 一致", "简化清单已写", "主要部件质量未漏装"],
    "A1-2": ["总质量误差 <5%", "Mass_Summary.csv 已输出", "质心对比图已保存"],
    "A2-1": ["θ₀ 参数 6–8 个且有 Bounds", "Joint_Theta0.csv 含来源说明", "非全刚性连接"],
    "A2-2": ["地脚 BC 与 SSS 一致", "重力方向正确", "Boundary_Spec.pdf 已签字"],
    "A3-1": ["粗/中/细三套网格已对比", "Δf₁₋₃ <2%", "结合部区域足够细"],
    "A4-1": ["静力变形趋势合理", "单位无混用", "Static 结果已存档"],
    "A4-2": ["模态基于预应力态", "设置参数已记录", "未与普通模态混用"],
    "A4-3": ["12 阶已提取", "6 阶弹性模态筛选完成", "剔除原因已记录"],
    "A5-1": ["P1–P12/E/TP 全部映射", "d>10 mm 已复核", "测点布置图已保存"],
    "A5-2": ["灵敏度分析完成", "试修参数优先级明确", "龙卷风图已保存"],
    "A6-1": ["FE_Y_Extreme 轴位与 B2 一致", "拖链/电缆姿态已核对", "与 Baseline 分开存档"],
    "A7-1": ["Baseline_FE_v0 已冻结", "Gate_G1 已签字", "通过后未改几何/网格/BC"],
    "B0-1": ["SSS_Checklist 逐项打勾", "四向照片 + 轴位截图齐全", "状态与 FE 一致"],
    "B0-2": ["温湿度已记录", "静置时间足够", "干扰源已注明"],
    "B0-3": ["背景噪底 ≥30 s", "峰值高于背景 20 dB", "噪底谱图已保存"],
    "B1-1": ["System_Config.txt 已保存", "试扫 SNR 合格", "Setup 现场照已拍"],
    "B1-2": ["区 A polygon 不含 P10–P12", "ScanZone_A.csv 已导出", "示意图已标注盲验区"],
    "B1-3": ["3 激励 ×5 trial 完成", "每次 Trial_Log 已填", "raw 入 E1_C0_SLDV_Raw/"],
    "B1-4": ["轻/重敲线性检验完成", "Linearity_Check.pdf 已存档", "正式测量力幅在线性区"],
    "B1-5": ["EMA P1–P9 完成", "P1–P3 同点 SLDV 已测", "E3_C0_EMA/ 已备份"],
    "B1-6": ["双备份当日完成", "备份路径写入矩阵", "原始文件夹设只读"],
    "B2-1": ["当日 B0-1~B0-3 已完成（SSS-H）", "Trial_Log_B2 + 双备份齐全", "未用 B2 数据 refit θ*"],
    "B3-1": ["当日 B0-1~B0-3 已完成", "修正未使用盲验数据", "Blind 分文件夹封存"],
    "B3-2": ["P10/P11/P12 各 5 重复", "Trial_Log_B3 + 双备份齐全", "三点空间分散"],
    "B3-3": ["Seal_Record 已填写", "B3 阶段未填 Unseal", "见证签字齐全"],
    "B4-1": ["仅重测失败步骤", "新编号不覆盖旧 raw", "重测原因已记录"],
    "C1": ["稳态图判据与 Settings 一致", "6 阶弹性模态已输出", "剔除极点有记录"],
    "C2": ["配对综合频率+MAC+形态", "歧义已记录", "Mode_Pairing_Table 已冻结前复核"],
    "C3": ["5 次重复统计完成", "Repeatability.csv 已输出", "误差棒图已保存"],
    "C4": ["Gate_G2 已签字", "配对表经第二人复核", "通过后未随意改配对"],
    "D1": ["Before 对比图已出", "修正必要性已量化", "配对错已排除"],
    "D2": ["J(θ) 仅用 P1–P9", "Objective_Function.txt 已写", "盲验与全场未进目标函数"],
    "D3": ["从 4 参数试修", "Iteration_Log 完整", "L-curve 已用于选参"],
    "D4": ["≥2 组初值已验证", "10start/ 已存档", "无非物理解未处理"],
    "D5": ["θ* 已收敛", "Updated_FE_ThetaStar 已导出", "Gate_G3 已签字"],
    "E1": ["Unseal_Record 日期正确", "θ* 确认冻结", "见证签字齐全"],
    "E2": ["三点 Updated/Baseline/Exp 齐全", "法向 MAC 已算", "Blind_Summary.csv 已出"],
    "E3": ["≥2/3 判定（MAC+频率双达标）", "平均 MAC±std 已报", "fail 未用于调参"],
    "E4": ["Gate_G4 已签字", "fail 有分析记录", "回 D3 未用盲验数据"],
    "F1": ["FE_Y_Extreme(θ*) 正推完成", "C1_Forward.csv 已出", "未 refit θ*"],
    "F2": ["与 B2 raw 已对比", "Limitation 已写", "表述与误差匹配"],
    "G1-1": ["Updated/Baseline 同频段", "Harmonic csv 已出", "基于 θ* 版本"],
    "G1-2": ["TP FRF 已输出", "模态叠加收敛已检", "Mode_Convergence.csv 已存"],
    "G2-1": ["附加质量位置已记录", "Added_Mass_Note 已写", "布置照已拍"],
    "G2-2": ["Harmonic_Excitation_Spec 已定", "力链标定完成", "输入力谱已测"],
    "G2-3": ["5 重复完成", "γ²≥0.7 已核查", "TP_FRF_5trials.csv 已出"],
    "G2-4": ["峰值频率为主指标", "Gate_G5 已签字", "三线 FRF 图已出"],
    "H1": ["创新点 3 条定稿", "Limitation 已写", "技术路线图已入第 1 章"],
    "H2": ["Figure_Index 已维护", "核心图 ≥16", "图号与数据路径一致"],
    "H3": ["Limitation_Checklist 6 条齐", "附录 A–D 同步", "结论与门禁一致"],
    "H4": ["符号表与正文一致", "参考文献 GB/T 7714", "无符号多义"],
    "H5": ["Public 版 FE/图已完成", "涉密信息已脱敏", "送审版可独立提交"],
    "H6": ["SCI 章节划分已规划", "避免自我抄袭", "扩展数据已在 I 阶段收齐"],
    "I1": ["04_SCI_Reserve/ 已建", "B1/B2/10-start 已隔离存档", "版本与 θ* 对应清晰"],
    "I2": ["5 trial raw 全保留", "02_Test_Raw/ 无删 trial", "processed 与 raw 分开"],
    "I3": ["EMA/C1/盲验汇总数据齐", "硕士论文可精简表述", "原始数据未补测缺口"],
}


def render_phase(doc: Document, phase_title: str, content) -> None:
    doc.add_paragraph(phase_title, style="Heading 1")
    if isinstance(content, list):
        for item in content:
            if len(item) == 3:
                code = item[0]
                add_step(
                    doc,
                    code,
                    item[1],
                    item[2],
                    CHECKPOINTS.get(code),
                    DELIVERABLES.get(code),
                )
    elif isinstance(content, dict):
        for sub_title, steps in content.items():
            doc.add_paragraph(sub_title, style="Heading 2")
            for item in steps:
                code = item[0]
                add_step(
                    doc,
                    code,
                    item[1],
                    item[2],
                    CHECKPOINTS.get(code),
                    DELIVERABLES.get(code),
                )


def build_document() -> Document:
    doc = Document()
    set_doc_style(doc)
    doc.add_paragraph("直驱三轴龙门动态模型确认", style="Title")
    add_para(
        doc,
        "逐步执行总表（文字详述版 · 硕士优先 · 数据一次收全 · 9+3 盲验）",
    )
    add_para(
        doc,
        """本文档按执行顺序列出从实验前准备到论文撰写的全部步骤。英文术语在中文后用括号备注；不写工期，只写步骤顺序。每一步结构为：文字说明 → 【必存数据】（编号列表，含文件名与用途）→ 【必存图片】（编号列表，含拍摄/导出要求）→ 现场操作要点。非必要不使用表格，便于开题后打印携带与实验现场查阅。""",
    )
    add_para(
        doc,
        """全文逻辑链为：阶段 0 定规则 → 阶段 A 建 Baseline FE 并通过 G1 → 阶段 B 一次收全实验 raw（B1 标定区 A，B2 Y 极限 raw，B3 盲验封存）→ 阶段 C 识别配对并通过 G2 → 阶段 D 仅用 P1–P9 修正得 θ* 并通过 G3 → 阶段 E 开封盲验并通过 G4 → 阶段 F 可选 C1 构型正推 → 阶段 G 谐响应并通过 G5 → 阶段 H 写论文 → 阶段 I SCI 数据储备。""",
    )

    order = [
        "阶段0：实验前准备（必须先于一切实验）",
        "阶段A：Baseline 仿真（基准有限元）",
        "阶段B：实验一次收全",
        "阶段C：实验模态识别与配对",
        "阶段D：模型修正（仅用 P1–P9）",
        "阶段E：盲验开封与判定",
        "阶段F：C1 构型正推验证",
        "阶段G：刀位谐响应仿真与实验",
        "阶段H：论文撰写",
        "阶段I：SCI 数据储备（硕士可少写、数据必须先收）",
    ]
    for key in order:
        render_phase(doc, key, STEPS[key])

    add_criteria_section(doc)
    add_fatal_section(doc)
    add_code_legend(doc)
    return doc


def main() -> None:
    doc = build_document()
    for path in OUTPUT_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print(f"Saved: {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
