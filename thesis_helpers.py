"""Thesis placeholder markers and docx helpers."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


def data(desc: str) -> str:
    return f"【待填数据：{desc}】"


def fig(label: str, caption: str) -> str:
    return f"【待插图：{label} {caption}】"


def tbl(label: str, caption: str) -> str:
    return f"【待填表：{label} {caption}】"


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)


def P(doc: Document, text: str, style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def H1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def H2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def H3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
