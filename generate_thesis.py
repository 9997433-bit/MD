#!/usr/bin/env python3
"""Render thesis_content to Word and run two-pass self-check."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from thesis_content import build_thesis_sections
from thesis_helpers import H1, H2, H3, P, bullets, setup_doc

OUTPUT = Path("/workspace/05_Thesis/直驱三轴龙门动态模型确认_论文初稿.docx")
OUTPUT_ALT = Path("/workspace/download/直驱三轴龙门动态模型确认_论文初稿.docx")
ZJU = Path("/workspace/浙大/毕业设计/仿真-实验大纲/直驱三轴龙门动态模型确认_论文初稿.docx")

REQUIRED_H1 = [
    "摘要",
    "Abstract",
    "第1章 绪论",
    "第2章 直驱龙门机床动态建模与模型确认的理论基础",
    "第3章 直驱三轴龙门机床有限元建模与基准仿真",
    "第4章 实验模态测试、模态识别与模型修正",
    "第5章 盲验、构型正推与刀位谐响应验证",
    "第6章 总结与展望",
    "参考文献",
    "附录A 坐标变换定义",
    "符号说明",
]

REQUIRED_TERMS = [
    "P1–P9",
    "P10–P12",
    "EX1",
    "SSS-M",
    "SSS-H",
    "G1",
    "G5",
    "θ*",
    "Updated",
    "Baseline",
    "【待填数据：",
    "【待插图：",
    "【待填表：",
]


def render_doc() -> Document:
    doc = Document()
    setup_doc(doc)
    for level, title, paras in build_thesis_sections():
        if level == "title":
            doc.add_paragraph(title, style="Title")
            continue
        if level == "h1":
            H1(doc, title)
            for para in paras:
                P(doc, para)
            continue
        if level == "h2":
            H2(doc, title)
            for para in paras:
                P(doc, para)
            continue
        if level == "h3":
            H3(doc, title)
            for para in paras:
                P(doc, para)
            continue
        if level == "bullets":
            bullets(doc, paras)
            continue
        if level == "p":
            for para in paras:
                P(doc, para)
    return doc


def extract_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def self_check_pass(name: str, text: str) -> list[str]:
    issues: list[str] = []
    # chapter headings
    for ch in REQUIRED_H1:
        if ch not in text:
            issues.append(f"缺少一级标题/区块：{ch}")
    # placeholders
    nd = len(re.findall(r"【待填数据：", text))
    nf = len(re.findall(r"【待插图：", text))
    nt = len(re.findall(r"【待填表：", text))
    if nd < 50:
        issues.append(f"待填数据占位过少：{nd}（期望≥50）")
    if nf < 16:
        issues.append(f"待插图占位过少：{nf}（期望≥16）")
    if nt < 10:
        issues.append(f"待填表占位过少：{nt}（期望≥10）")
    # terminology
    for term in REQUIRED_TERMS:
        if term not in text:
            issues.append(f"缺少关键术语/标记：{term}")
    # no unfilled numeric claims like "为 12.3 Hz" without placeholder - heuristic
    bad_nums = re.findall(r"(?<![待填数据：\d])[\d]+\.[\d]+\s*Hz", text)
    if bad_nums:
        issues.append(f"可能存在未占位数值（Hz）：{bad_nums[:5]}")
    # chapter 6 limitation
    if "Limitation" not in text and "不足" not in text:
        issues.append("第6章 Limitation/不足 表述可能缺失")
    # blind validation
    if "盲验" not in text or "Pre-registration" not in text:
        issues.append("盲验/Pre-registration 表述缺失")
    print(f"--- 自检 {name} ---")
    print(f"  字符数: {len(text)}")
    print(f"  待填数据: {nd}  待插图: {nf}  待填表: {nt}")
    if issues:
        print(f"  问题 {len(issues)} 项:")
        for i in issues:
            print(f"    - {i}")
    else:
        print("  通过")
    return issues


def main() -> None:
    doc = render_doc()
    text = extract_text(doc)
    issues1 = self_check_pass("第1轮", text)
    issues2 = self_check_pass("第2轮", text)
    if issues1 or issues2:
        # second pass should be identical if no mutation; log but still save draft
        print("警告：自检发现问题，已仍保存初稿供人工回填。")
    for path in (OUTPUT, OUTPUT_ALT, ZJU):
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print(f"Saved: {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
